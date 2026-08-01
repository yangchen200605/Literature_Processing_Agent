"""Word / PDF 解析（文本+表格+图片元信息+封面）与摘要导出。"""

from __future__ import annotations

import base64
import io
import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fastapi import HTTPException, UploadFile
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.storage import save_upload

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_COVER_IMAGES = 6

pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [(r + [""] * width)[:width] for r in rows]
    header = norm[0]
    lines = [
        "| " + " | ".join(c.replace("\n", " ").strip() or " " for c in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in norm[1:]:
        lines.append("| " + " | ".join(c.replace("\n", " ").strip() or " " for c in row) + " |")
    return "\n".join(lines)


async def parse_document_upload(file: UploadFile) -> dict:
    """解析上传文档：提取文本/表格、生成封面、保存原件供下载。"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="未提供文件名")

    suffix = Path(file.filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="仅支持 PDF (.pdf) 或 Word (.docx) 文件")
    if suffix == ".doc":
        raise HTTPException(status_code=400, detail="暂不支持旧版 .doc，请另存为 .docx 后上传")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="文件为空")
    if len(data) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="文件过大，请控制在 20MB 以内")

    try:
        if suffix == ".pdf":
            parsed = _parse_pdf(data)
        else:
            parsed = _parse_docx(data)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"无法解析文件: {e}") from e

    text = (parsed.get("text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="未能从文件中提取到文本内容")

    cover_png: bytes | None = parsed.get("cover_png")
    file_id = save_upload(
        filename=file.filename,
        file_bytes=data,
        cover_png=cover_png,
        text=text,
        meta_extra={
            "page_count": parsed.get("page_count", 0),
            "table_count": parsed.get("table_count", 0),
            "image_count": parsed.get("image_count", 0),
            "file_type": suffix.lstrip("."),
        },
    )

    cover_data_url = None
    if cover_png:
        cover_data_url = "data:image/png;base64," + base64.b64encode(cover_png).decode("ascii")

    previews = []
    for item in parsed.get("image_previews") or []:
        previews.append(
            {
                "label": item["label"],
                "data_url": "data:image/png;base64,"
                + base64.b64encode(item["png"]).decode("ascii"),
            }
        )

    return {
        "file_id": file_id,
        "filename": file.filename,
        "file_type": suffix.lstrip("."),
        "text": text,
        "cover_url": f"/api/uploads/{file_id}/cover" if cover_png else None,
        "cover_data_url": cover_data_url,
        "download_url": f"/api/uploads/{file_id}/download",
        "page_count": parsed.get("page_count", 0),
        "table_count": parsed.get("table_count", 0),
        "image_count": parsed.get("image_count", 0),
        "char_count": len(text),
        "image_previews": previews,
    }


# 兼容旧调用名
async def extract_text_from_upload(file: UploadFile) -> tuple[str, str]:
    result = await parse_document_upload(file)
    return result["text"], result["filename"]


def _pixmap_to_preview_png(pix: fitz.Pixmap, max_side: int = 480) -> bytes | None:
    try:
        if pix.n - pix.alpha >= 4:  # CMYK etc.
            pix = fitz.Pixmap(fitz.csRGB, pix)
        png = pix.tobytes("png")
        if max(pix.width, pix.height) <= max_side and len(png) < 1_500_000:
            return png
        src = fitz.open(stream=png, filetype="png")
        try:
            page = src[0]
            scale = min(1.0, max_side / max(page.rect.width, page.rect.height, 1))
            thumb = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            out = thumb.tobytes("png")
            return out if len(out) < 1_500_000 else None
        finally:
            src.close()
    except Exception:
        return None


def _parse_pdf(data: bytes) -> dict:
    parts: list[str] = []
    table_count = 0
    image_count = 0
    image_previews: list[dict] = []
    cover_png: bytes | None = None
    page_count = 0

    with fitz.open(stream=data, filetype="pdf") as doc:
        page_count = doc.page_count
        if page_count:
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            cover_png = pix.tobytes("png")

        seen_xrefs: set[int] = set()
        for page_index, page in enumerate(doc):
            page_bits: list[str] = []
            text = page.get_text("text").strip()
            if text:
                page_bits.append(text)

            try:
                finder = page.find_tables()
                tables = list(finder.tables) if finder else []
            except Exception:
                tables = []

            for table in tables:
                try:
                    rows = table.extract()
                except Exception:
                    continue
                clean_rows = [
                    [str(c or "").strip() for c in row]
                    for row in (rows or [])
                    if any(str(c or "").strip() for c in row)
                ]
                if not clean_rows:
                    continue
                table_count += 1
                md = _table_to_markdown(clean_rows)
                page_bits.append(f"\n[表格 {table_count} · 第{page_index + 1}页]\n{md}\n")

            for img in page.get_images(full=True):
                xref = img[0]
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)
                image_count += 1
                page_bits.append(f"[图 {image_count} · 第{page_index + 1}页]")
                if len(image_previews) < MAX_COVER_IMAGES:
                    try:
                        pix = fitz.Pixmap(doc, xref)
                        png_bytes = _pixmap_to_preview_png(pix)
                        if png_bytes:
                            image_previews.append(
                                {
                                    "label": f"图{image_count} · P{page_index + 1}",
                                    "png": png_bytes,
                                }
                            )
                    except Exception:
                        pass

            if page_bits:
                parts.append(f"## 第 {page_index + 1} 页\n\n" + "\n\n".join(page_bits))

    return {
        "text": "\n\n".join(parts),
        "cover_png": cover_png,
        "page_count": page_count,
        "table_count": table_count,
        "image_count": image_count,
        "image_previews": image_previews,
    }


def _parse_docx(data: bytes) -> dict:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    table_count = 0
    image_count = 0
    image_previews: list[dict] = []
    cover_png: bytes | None = None

    # 按文档顺序遍历段落与表格
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # 找段落文字
            texts = [t.text for t in child.iter(qn("w:t")) if t.text]
            para = "".join(texts).strip()
            # 段落内图片
            drawings = list(child.iter(qn("w:drawing")))
            if drawings:
                for _ in drawings:
                    image_count += 1
                    parts.append(f"[图 {image_count}]")
            if para:
                parts.append(para)
        elif tag == "tbl":
            rows: list[list[str]] = []
            for tr in child.iter(qn("w:tr")):
                cells: list[str] = []
                for tc in tr.iter(qn("w:tc")):
                    cell_text = "".join(
                        t.text for t in tc.iter(qn("w:t")) if t.text
                    ).strip()
                    cells.append(cell_text)
                if any(cells):
                    rows.append(cells)
            if rows:
                table_count += 1
                parts.append(f"\n[表格 {table_count}]\n{_table_to_markdown(rows)}\n")

    # 提取嵌入图片作封面/预览
    for rel in document.part.rels.values():
        if "image" not in getattr(rel, "reltype", ""):
            continue
        try:
            blob = rel.target_part.blob
        except Exception:
            continue
        if not blob:
            continue
        # 已在正文计数过 drawing；这里补预览（避免重复计数只用于预览）
        try:
            # 统一转 png 缩略图
            img_doc = fitz.open(stream=blob, filetype="png")
        except Exception:
            try:
                img_doc = fitz.open(stream=blob, filetype="jpeg")
            except Exception:
                try:
                    img_doc = fitz.open(stream=blob)
                except Exception:
                    continue
        try:
            page = img_doc[0]
            scale = min(1.0, 480 / max(page.rect.width, page.rect.height, 1))
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            png_bytes = pix.tobytes("png")
            if cover_png is None:
                cover_png = png_bytes
            if len(image_previews) < MAX_COVER_IMAGES and len(png_bytes) < 1_500_000:
                image_previews.append(
                    {
                        "label": f"图{len(image_previews) + 1}",
                        "png": png_bytes,
                    }
                )
        finally:
            img_doc.close()

    # 若没有图，生成简易封面（文件名文字页）
    if cover_png is None:
        cover_png = _render_text_cover("Word 文档", "")

    # 若正文没扫到 drawing 但 rel 有图，修正 image_count
    if image_count == 0 and image_previews:
        image_count = len(image_previews)
        for i in range(1, image_count + 1):
            if f"[图 {i}]" not in "\n".join(parts):
                parts.append(f"[图 {i}]")

    text = "\n\n".join(p for p in parts if p and str(p).strip())
    return {
        "text": text,
        "cover_png": cover_png,
        "page_count": 1,
        "table_count": table_count,
        "image_count": image_count,
        "image_previews": image_previews,
    }


def _render_text_cover(title: str, subtitle: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=420, height=594)
    page.draw_rect(page.rect, color=(0.93, 0.94, 0.97), fill=(0.93, 0.94, 0.97))
    page.insert_textbox(
        fitz.Rect(40, 200, 380, 280),
        title[:80],
        fontsize=22,
        color=(0.15, 0.2, 0.35),
        align=1,
    )
    if subtitle:
        page.insert_textbox(
            fitz.Rect(40, 290, 380, 340),
            subtitle[:120],
            fontsize=12,
            color=(0.4, 0.45, 0.55),
            align=1,
        )
    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
    png = pix.tobytes("png")
    doc.close()
    return png


def export_docx(content: str) -> bytes:
    """将摘要 Markdown/纯文本导出为 Word。"""
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    title = document.add_heading("文献摘要", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("### "):
            document.add_heading(_plain_md(stripped[4:].strip()), level=3)
        elif stripped.startswith("## "):
            document.add_heading(_plain_md(stripped[3:].strip()), level=2)
        elif stripped.startswith("# "):
            document.add_heading(_plain_md(stripped[2:].strip()), level=1)
        elif stripped.startswith(("- ", "* ")):
            p = document.add_paragraph(style="List Bullet")
            _apply_bold_runs(p, stripped[2:].strip())
        else:
            p = document.add_paragraph()
            _apply_bold_runs(p, stripped)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _plain_md(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _apply_bold_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def export_pdf(content: str) -> bytes:
    """将摘要导出为 PDF（支持中文）。"""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CNBody",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
    )
    h1 = ParagraphStyle(
        "CNH1",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=22,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "CNH2",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=13,
        leading=20,
        spaceAfter=8,
    )

    story = [Paragraph("文献摘要", h1), Spacer(1, 8)]
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        safe = (
            stripped.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)

        if stripped.startswith("### "):
            story.append(Paragraph(safe[4:].strip(), h2))
        elif stripped.startswith("## "):
            story.append(Paragraph(safe[3:].strip(), h2))
        elif stripped.startswith("# "):
            story.append(Paragraph(safe[2:].strip(), h1))
        elif stripped.startswith(("- ", "* ")):
            story.append(Paragraph(f"• {safe[2:].strip()}", body))
        else:
            story.append(Paragraph(safe, body))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()
